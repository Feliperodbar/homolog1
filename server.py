import base64
import io
import time
import os
from flask import Flask, send_from_directory, request, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Inches

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)


@app.route('/')
def index():
    return send_from_directory('.', 'index.html')


@app.route('/<path:path>')
def static_proxy(path):
    return send_from_directory('.', path)


@app.route('/health')
def health():
    return {'status': 'ok'}


@app.route('/export-docx', methods=['POST'])
def export_docx():
    data = request.get_json(silent=True) or {}
    steps = data.get('steps', [])

    doc = Document()
    doc.add_heading('Guia Homolog', level=1)

    for i, s in enumerate(steps, 1):
        title = s.get('title') or 'Passo'
        doc.add_heading(f"{i}. {title}", level=2)

        tag = s.get('tag')
        if tag:
            p = doc.add_paragraph()
            run = p.add_run(f"Tag: {tag}")
            run.bold = True

        desc = s.get('description')
        if desc:
            doc.add_paragraph(desc)

        img_data_url = s.get('imageDataUrl')
        if img_data_url:
            try:
                header, b64 = img_data_url.split(',', 1)
                img_bytes = base64.b64decode(b64)
                bio = io.BytesIO(img_bytes)
                bio.seek(0)
                # Largura ~6.5" para caber na página
                doc.add_picture(bio, width=Inches(6.5))
            except Exception:
                doc.add_paragraph('[Imagem indisponível]')

        doc.add_paragraph('')

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    filename = f"homolog_export_{int(time.time())}.docx"
    return send_file(
        out,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename,
    )


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8010))
    app.run(host='0.0.0.0', port=port)
